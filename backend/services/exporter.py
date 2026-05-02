import io
import html
import os
import json
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from services import license_service
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, NextPageTemplate, PageBreak, Flowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY, TA_CENTER, TA_RIGHT, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import Color, gray

# --- BETA PHASE PROTECTION ---
# Watermarking is active by default for all unactivated drafts
def get_protection_status():
    return not license_service.is_activated()

BETA_LABEL = "Tome-Master BETA - UNSUBMITTED DRAFT"
# -----------------------------

def add_toc_field(paragraph):
    """Injects a dynamic Microsoft Word Table of Contents field code."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2); run._r.append(fldChar3)

RED_SPELL = RGBColor(0xEF, 0x44, 0x44); BLUE_GRAMMAR = RGBColor(0x3B, 0x82, 0xF6)

def _add_inline_runs(p_obj, node, is_bold=False, is_italic=False):
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if not text: continue
            run = p_obj.add_run(text); run.bold = is_bold; run.italic = is_italic
        elif isinstance(child, Tag):
            classes = child.get('class') or []
            if 'misspelled-word' in classes:
                run = p_obj.add_run(child.get_text()); run.bold = is_bold; run.italic = is_italic; run.underline = True; run.font.color.rgb = RED_SPELL
            elif 'grammar-squiggle' in classes:
                run = p_obj.add_run(child.get_text()); run.bold = is_bold; run.italic = is_italic; run.underline = True; run.font.color.rgb = BLUE_GRAMMAR
            elif child.name in ('strong', 'b'): _add_inline_runs(p_obj, child, is_bold=True, is_italic=is_italic)
            elif child.name in ('em', 'i'): _add_inline_runs(p_obj, child, is_bold=is_bold, is_italic=True)
            else: _add_inline_runs(p_obj, child, is_bold=is_bold, is_italic=is_italic)

def _strip_toc(soup):
    toc_div = soup.find('div', class_='editor-toc')
    if toc_div: toc_div.decompose()
    
    found_manual = False
    for tag in soup.find_all(['h1', 'h2', 'h3', 'p', 'b', 'strong']):
        if not tag.parent: continue
        text = tag.get_text().replace('\u200b', '').strip().lower()
        if text in ["table of contents", "contents"]:
            found_manual = True
            placeholder = soup.new_tag('div', id='toc-placeholder')
            tag.insert_before(placeholder)
            
            curr = tag.next_sibling
            while curr:
                next_node = curr.next_sibling
                if getattr(curr, 'name', None) in ['h1', 'h2', 'h3'] and curr.get_text().strip(): break
                curr_text = curr.get_text().replace('\u200b', '').strip().lower() if hasattr(curr, 'get_text') else ""
                if curr_text and (curr_text.startswith('chapter ') or curr_text in ['prologue', 'epilogue']): break
                if hasattr(curr, 'decompose'): curr.decompose()
                curr = next_node
            tag.decompose()
            break
            
    if not found_manual:
        for tag in soup.find_all(['h1', 'h2', 'h3']):
            t = tag.get_text().replace('\u200b', '').lower()
            if any(x in t for x in ['chapter', 'prologue', 'part 1', 'epilogue']):
                placeholder = soup.new_tag('div', id='toc-placeholder')
                tag.insert_before(placeholder)
                found_manual = True
                break
        if not found_manual:
            placeholder = soup.new_tag('div', id='toc-placeholder')
            soup.insert(0, placeholder)

def _extract_frontmatter(soup, default_title, default_author):
    title = default_title
    author = default_author
    
    blocks = soup.find_all(['h1', 'h2', 'h3', 'p', 'div'])
    
    found_title = False
    for node in blocks[:20]:
        if not getattr(node, 'name', None): continue
            
        text = node.get_text().replace('\u200b', '').replace('\xa0', ' ').strip()
        if not text: continue
            
        if not found_title:
            title = text
            found_title = True
            node.decompose()
        else:
            t_lower = text.lower()
            if t_lower in ['by', 'by:', 'by-'] or t_lower.startswith('by ') or node.name in ['h2', 'h3'] or len(text.split()) <= 5:
                author = text
                curr = node.next_sibling # Get pointer BEFORE destroying the node
                node.decompose()
                
                if t_lower in ['by', 'by:', 'by-']:
                    while curr:
                        next_curr = curr.next_sibling
                        if getattr(curr, 'name', None) in ['p', 'h1', 'h2', 'h3', 'div']:
                            n_text = curr.get_text().replace('\u200b', '').replace('\xa0', ' ').strip()
                            if n_text:
                                author += " " + n_text
                                curr.decompose()
                                break
                            else:
                                curr.decompose()
                        curr = next_curr
            break
            
    return title, author

def _add_docx_beta_watermark(doc):
    """Adds a prominent Beta watermark disclaimer to the header of every section."""
    if not get_protection_status(): return
    for section in doc.sections:
        header = section.header
        # Check if header already has text to avoid double-adding
        if not header.paragraphs:
            p = header.add_paragraph()
        else:
            p = header.paragraphs[0]
        p.text = f"--- {BETA_LABEL} ---"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 128, 128)

def generate_docx(content: str, chapters: list = None, title: str = "Manuscript Title", author: str = "Author Name", output_format: str = "chicago", cover_image: str = None) -> io.BytesIO:
    doc = Document()
    _add_docx_beta_watermark(doc)
    
    if cover_image:
        try:
            import base64
            header, encoded = cover_image.split(",", 1)
            img_data = base64.b64decode(encoded)
            img_stream = io.BytesIO(img_data)
            doc.add_picture(img_stream, width=Inches(6))
            doc.add_page_break()
        except:
            pass # Silently fail if image data is corrupted
            
    for section in doc.sections:
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    
    style_normal = doc.styles['Normal']; font = style_normal.font
    font.name = 'Garamond' if output_format == 'penguin' else 'Times New Roman'
    font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.5 if output_format == 'penguin' else 2.0
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    for i in range(1, 4):
        style_name = f'TOC {i}'
        if style_name in doc.styles:
            s = doc.styles[style_name]; s.font.name = font.name; s.font.size = Pt(12); s.paragraph_format.line_spacing = style_normal.paragraph_format.line_spacing
            
    soup = BeautifulSoup(content, 'html.parser')
    _strip_toc(soup)
    
    # We no longer extract and destroy frontmatter because the user wants the title to print exactly as formatted in the DOM
    # title, author = _extract_frontmatter(soup, title, author)

    has_content = False
    for node in soup.find_all(['h1', 'h2', 'h3', 'p', 'div']):
        if node.name == 'div' and node.get('id') == 'toc-placeholder':
            if has_content: doc.add_page_break()
            doc.add_heading('Table of Contents', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("Select this field and press F9 or right-click 'Update Field' to generate the Table of Contents.")
            add_toc_field(doc.add_paragraph())
            has_content = True
            continue
            
        text = node.get_text().replace('\u200b', '').replace('\xa0', ' ').strip()
        if not text: continue
        
        if node.name in ['h1', 'h2', 'h3']:
            if has_content: doc.add_page_break()
            doc.add_heading(text, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            has_content = True
        elif node.name == 'p':
            if text in ['***', '* * *', '#']:
                doc.add_paragraph('#').alignment = WD_ALIGN_PARAGRAPH.CENTER
                has_content = True
                continue
            p_obj = doc.add_paragraph(); p_obj.paragraph_format.first_line_indent = Inches(0.5); _add_inline_runs(p_obj, node)
            has_content = True
    
    file_stream = io.BytesIO(); doc.save(file_stream); file_stream.seek(0); return file_stream

def generate_analysis_docx(markdown_text: str) -> io.BytesIO:
    doc = Document()
    _add_docx_beta_watermark(doc)
    for s in doc.sections: s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)
    doc.add_heading("tome_master AI Boardroom Report", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1; continue
        if line.startswith('# '): doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '): doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '): doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith(('- ', '* ')):
            p = doc.add_paragraph(style='List Bullet'); p.add_run(line[2:].strip().replace('**', ''))
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) > 2:
                # Filter out empty strings from split and ignore decorative separator lines
                headers = [val.strip() for val in table_lines[0].split('|') if val.strip()]
                if not headers: continue
                
                table = doc.add_table(rows=1, cols=len(headers), style='Table Grid')
                hdr_cells = table.rows[0].cells
                for idx, hdr in enumerate(headers):
                    hdr_cells[idx].text = hdr
                
                # Make headers repeat on each page for context
                tr = table.rows[0]._tr
                trPr = tr.get_or_add_trPr()
                tblHeader = OxmlElement('w:tblHeader')
                trPr.append(tblHeader)
                
                for r_idx, row_line in enumerate(table_lines):
                    # Skip header line (0) and alignment separator lines (1)
                    if r_idx <= 1:
                        continue
                    
                    row_data = [val.strip() for val in row_line.split('|') if val.strip()]
                    if len(row_data) > 0:
                        row_cells = table.add_row().cells
                        # Ensure we don't exceed header count but fill what we have
                        for idx, val in enumerate(row_data[:len(headers)]):
                            row_cells[idx].text = val
            continue
        else:
            p = doc.add_paragraph(); p.add_run(line.replace('**', '')); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        i += 1
        
    file_stream = io.BytesIO(); doc.save(file_stream); file_stream.seek(0); return file_stream

TRADE_PAPERBACK = (6 * inch, 9 * inch)

def to_roman(num):
    lookup = [(1000,'m'),(900,'cm'),(500,'d'),(400,'cd'),(100,'c'),(90,'xc'),(50,'l'),(40,'xl'),(10,'x'),(9,'ix'),(5,'v'),(4,'iv'),(1,'i')]
    r = ''; 
    for v, n in lookup:
        while num >= v: r += n; num -= v
    return r or 'i'

class ResetPageFlowable(Flowable):
    def wrap(self, w, h): return (0, 0)
    def draw(self): self.canv._pageNumber = 0

def draw_watermark(canvas, doc):
    if not get_protection_status(): return
    canvas.saveState()
    canvas.setFont('Times-Bold', 60)
    canvas.setFillAlpha(0.15)
    canvas.setStrokeColor(gray)
    canvas.translate(TRADE_PAPERBACK[0]/2, TRADE_PAPERBACK[1]/2)
    canvas.rotate(45)
    canvas.drawCentredString(0, 0, BETA_LABEL)
    canvas.restoreState()

def generate_pdf(content: str, chapters: list = None, title: str = "Manuscript Title", author: str = "Author Name", output_format: str = "chicago", cover_image: str = None) -> io.BytesIO:
    file_stream = io.BytesIO(); margin = 0.75 * inch; gutter = 0.25 * inch 
    frame = Frame(margin + gutter, margin, TRADE_PAPERBACK[0] - 2*margin - gutter, TRADE_PAPERBACK[1] - 2*margin, id='normal')
    
    from reportlab.platypus import Image
    
    def draw_f(canvas, doc):
        draw_watermark(canvas, doc)
        canvas.saveState(); canvas.setFont('Times-Roman', 10); canvas.drawCentredString(TRADE_PAPERBACK[0]/2.0, margin/2.0, to_roman(doc.page)); canvas.restoreState()
    def draw_n(canvas, doc):
        draw_watermark(canvas, doc)
        canvas.saveState(); canvas.setFont('Times-Roman', 10); canvas.drawCentredString(TRADE_PAPERBACK[0]/2.0, margin/2.0, str(doc.page)); canvas.restoreState()
    doc = BaseDocTemplate(file_stream, pagesize=TRADE_PAPERBACK)
    doc.addPageTemplates([PageTemplate(id='FrontMatter', frames=frame, onPage=draw_f), PageTemplate(id='Narrative', frames=frame, onPage=draw_n)])
    styles = getSampleStyleSheet()
    
    leading_val = 15 if output_format == "penguin" else 22
    style_n = ParagraphStyle('CM_N', parent=styles['Normal'], fontName='Times-Roman', fontSize=11, leading=leading_val, firstLineIndent=0.3*inch, alignment=TA_JUSTIFY)
    style_f = ParagraphStyle('CM_F', parent=style_n, firstLineIndent=0)
    style_t = ParagraphStyle('CM_T', parent=styles['Heading1'], fontName='Times-Bold', fontSize=24, leading=30, alignment=TA_CENTER, spaceAfter=inch)
    style_c = ParagraphStyle('CM_C', parent=styles['Heading2'], fontName='Times-Bold', fontSize=16, leading=20, alignment=TA_CENTER, spaceBefore=1.5*inch, spaceAfter=0.5*inch)
    
    style_toc = ParagraphStyle('TOC_Item', parent=style_n, firstLineIndent=0)
    
    story = []
    
    if cover_image:
        try:
            import base64
            header, encoded = cover_image.split(",", 1)
            img_data = base64.b64decode(encoded)
            img_stream = io.BytesIO(img_data)
            # Standard Trade Paperback 6x9 fitting
            img = Image(img_stream, width=4*inch, height=6*inch)
            story.append(img)
            story.append(PageBreak())
        except:
            pass

    soup = BeautifulSoup(content, 'html.parser')
    _strip_toc(soup)
    
    found_n = False
    has_content = False
    for i, node in enumerate(soup.find_all(['h1', 'h2', 'h3', 'p', 'div'])):
        if node.name == 'div' and node.get('id') == 'toc-placeholder':
            if chapters:
                if has_content: story.append(PageBreak())
                story.append(Paragraph("Table of Contents", style_c))
                for idx, chap in enumerate(chapters):
                    t = (chap.get('suggested_title') or chap.get('title') or "").replace('Chapter','').replace(':','').strip()
                    story.append(Paragraph(f'{html.escape(t)} &nbsp;&nbsp;&mdash;&nbsp;&nbsp; Page {chap.get("display_page") or 1}', style_toc))
                has_content = True
            continue
            
        text = node.get_text().replace('\u200b', '').replace('\xa0', ' ').strip()
        if not text: continue
        
        if node.name in ['h1', 'h2', 'h3']:
            if has_content:
                if not found_n and any(x in text.lower() for x in ['chapter', 'prologue', 'part 1']):
                    story.append(NextPageTemplate('Narrative')); story.append(PageBreak()); story.append(ResetPageFlowable()); found_n = True
                else: story.append(PageBreak())
            story.append(Paragraph(f'{html.escape(text)}', style_c))
            has_content = True
        elif node.name == 'p':
            if text in ['***', '* * *', '#']:
                story.append(Spacer(1, 0.3*inch)); story.append(Paragraph('#', ParagraphStyle('SB', parent=style_n, alignment=TA_CENTER))); story.append(Spacer(1, 0.3*inch))
                has_content = True
                continue
            story.append(Paragraph(_node_to_reportlab_markup(node), style_n))
            has_content = True
    doc.build(story); file_stream.seek(0); return file_stream

def _node_to_reportlab_markup(node) -> str:
    parts = []
    for child in node.children:
        if isinstance(child, NavigableString): parts.append(html.escape(str(child)))
        elif isinstance(child, Tag):
            classes = child.get('class') or []; inner = _node_to_reportlab_markup(child)
            if 'misspelled-word' in classes: parts.append(f'<u><font color="#ef4444">{inner}</font></u>')
            elif 'grammar-squiggle' in classes: parts.append(f'<u><font color="#3b82f6">{inner}</font></u>')
            elif child.name in ('strong', 'b'): parts.append(f'<b>{inner}</b>')
            elif child.name in ('em', 'i'): parts.append(f'<i>{inner}</i>')
            else: parts.append(inner)
    return ''.join(parts)

import tempfile
import uuid
from ebooklib import epub

def generate_epub(content: str, chapters: list = None, title: str = "Manuscript Title", author: str = "Author Name", output_format: str = "chicago", cover_image: str = None) -> io.BytesIO:
    book = epub.EpubBook()
    
    if cover_image:
        try:
            import base64
            header, encoded = cover_image.split(",", 1)
            img_data = base64.b64decode(encoded)
            book.set_cover("cover.jpg", img_data)
        except:
            pass

    book.set_identifier(str(uuid.uuid4()))
    book.set_title(title)
    book.set_language('en')
    book.add_author(author)
    
    style = 'body { font-family: "Times New Roman", Times, serif; line-height: 2.0; } h1, h2, h3 { text-align: center; } p { text-align: justify; text-indent: 1.5em; margin-top: 0; margin-bottom: 0; }'
    if output_format == "penguin":
        style = 'body { font-family: "Garamond", "Palatino Linotype", serif; line-height: 1.5; } h1, h2, h3 { text-align: center; } p { text-align: justify; text-indent: 1.5em; margin-top: 0; margin-bottom: 0; }'
    default_css = epub.EpubItem(uid="style_default", file_name="style/default.css", media_type="text/css", content=style)
    book.add_item(default_css)
    
    soup = BeautifulSoup(content, 'html.parser')
    _strip_toc(soup)
    
    sections = []
    current_chapter_title = "Front Matter"
    current_html_blocks = []
    
    if get_protection_status():
        current_html_blocks.append(f'<div style="border: 2px solid #666; padding: 2em; margin: 2em; text-align: center; border-radius: 10px;">')
        current_html_blocks.append(f'<h1 style="color: #666;">BETA ACCESS NOTICE</h1>')
        current_html_blocks.append(f'<p style="text-align: center; text-indent: 0;">This digital proof was generated during the <strong>tome_master Beta Phase</strong>.</p>')
        current_html_blocks.append(f'<p style="text-align: center; text-indent: 0; color: #ef4444; font-weight: bold;">{BETA_LABEL}</p>')
        current_html_blocks.append(f'<p style="text-align: center; text-indent: 0; font-size: 0.8em; margin-top: 1em;">Distribution or public sharing of this specific proof is discouraged to protect the development of the author\'s IP.</p>')
        current_html_blocks.append(f'</div>')
        current_html_blocks.append('<div style="page-break-after: always;"></div>')
    
    for tag in soup.find_all(['h1', 'h2', 'h3', 'p', 'div']):
        if tag.name == 'div' and tag.get('id') == 'toc-placeholder':
            continue 
            
        text = tag.get_text().replace('\u200b', '').replace('\xa0', ' ').strip()
        if not text: continue
        
        if tag.name in ['h1', 'h2', 'h3']:
            if current_html_blocks:
                sections.append((current_chapter_title, current_html_blocks))
            current_chapter_title = text
            current_html_blocks = [f"<h1>{html.escape(text)}</h1>"]
        elif tag.name == 'p':
            if text in ['***', '* * *', '#']:
                current_html_blocks.append('<p style="text-align: center;">***</p>')
            else:
                html_inner = "".join([str(c) for c in tag.contents])
                current_html_blocks.append(f"<p>{html_inner}</p>")
                
    if current_html_blocks:
        sections.append((current_chapter_title, current_html_blocks))
        
    epub_chapters = []
    for i, (chap_title, blocks) in enumerate(sections):
        c = epub.EpubHtml(title=chap_title, file_name=f'chap_{i}.xhtml', lang='en')
        c.content = f'<html><head><link rel="stylesheet" href="style/default.css" type="text/css"/></head><body>{"".join(blocks)}</body></html>'
        c.add_item(default_css)
        book.add_item(c)
        epub_chapters.append(c)
        
    book.toc = tuple(epub_chapters)
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    # The 'nav' item in the spine makes a visual Table of Contents page at the start of the book.
    # Users prefer to use the native TOC menu of the Ebook reader, so we remove 'nav' from the spine
    # but keep it in the manifest (by calling add_item above) as required by the EPUB3 spec.
    book.spine = epub_chapters
    
    with tempfile.NamedTemporaryFile(suffix='.epub', delete=False) as tmp:
        tmp_name = tmp.name
        
    epub.write_epub(tmp_name, book, {})
        
    out_stream = io.BytesIO()
    with open(tmp_name, 'rb') as f:
        out_stream.write(f.read())
    out_stream.seek(0)
    import os
    os.remove(tmp_name)
    
    return out_stream

